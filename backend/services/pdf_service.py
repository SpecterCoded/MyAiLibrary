import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader
import os
import shutil
from services.dependency_failure_service import local_path_failure, missing_configuration


def _get_tesseract_path(user_id: str | None = None) -> str:
    """Get Tesseract path from user settings or auto-detect."""
    # Try user settings first
    if user_id:
        try:
            from database import SessionLocal
            from models import UserSetting
            db = SessionLocal()
            try:
                settings = db.query(UserSetting).filter(UserSetting.user_id == user_id).first()
                if settings and settings.tesseract_path:
                    return settings.tesseract_path
            finally:
                db.close()
        except Exception:
            pass

    # Try PATH
    path = shutil.which("tesseract")
    if path:
        return path

    # Try common paths
    common_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for p in common_paths:
        if os.path.exists(p):
            return p

    return ""


def extract_pdf_text(file_path: str, user_id: str | None = None) -> str:
    """Extract text from PDF. Auto-detects scanned vs digital PDF."""
    reader = PdfReader(file_path)

    text = ""
    for i, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text:
            text += f"[Page {i}]\n{page_text}\n"

    # If we got meaningful text, it's a digital PDF
    if len(text.strip()) > 50:
        return text

    # Otherwise, it's likely a scanned PDF — use pytesseract on each page
    return _extract_with_ocr_pdf(file_path, user_id)



def _extract_with_ocr_pdf(file_path: str, user_id: str | None = None) -> str:
    """Extract text from scanned PDF using pytesseract."""
    import fitz  # PyMuPDF

    tesseract_path = _get_tesseract_path(user_id)
    if not tesseract_path:
        raise missing_configuration(service="Tesseract OCR", stage="indexing", settings_section="Tesseract OCR", fields=["Tesseract executable path"])
    if not os.path.isfile(tesseract_path):
        raise local_path_failure(code="path_not_found", service="Tesseract OCR", stage="indexing", settings_section="Tesseract OCR", path_label="Tesseract executable path")
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    doc = fitz.open(file_path)
    text = ""

    for i, page in enumerate(doc, 1):
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        page_text = pytesseract.image_to_string(img)
        if page_text.strip():
            text += f"[Page {i}]\n{page_text}\n"

    doc.close()
    return text


def extract_docx_text(file_path: str):
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text = []

    for paragraph in doc.paragraphs:
        paragraph_text = (paragraph.text or "").strip()
        if not paragraph_text:
            continue
        style_name = ""
        try:
            style_name = (paragraph.style.name or "").lower()
        except Exception:
            style_name = ""
        if style_name.startswith("heading"):
            text.append(f"# {paragraph_text}")
        elif paragraph.style and "list" in style_name:
            text.append(f"- {paragraph_text}")
        else:
            text.append(paragraph_text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [(cell.text or "").strip() for cell in row.cells]
            values = [value for value in values if value]
            if values:
                rows.append(" | ".join(values))
        if rows:
            text.append("[Table]")
            text.extend(rows)

    return "\n".join(text)


def extract_image_text(file_path: str, user_id: str | None = None):
    """Extract text from image using pytesseract."""
    tesseract_path = _get_tesseract_path(user_id)
    if not tesseract_path:
        raise missing_configuration(service="Tesseract OCR", stage="indexing", settings_section="Tesseract OCR", fields=["Tesseract executable path"])
    if not os.path.isfile(tesseract_path):
        raise local_path_failure(code="path_not_found", service="Tesseract OCR", stage="indexing", settings_section="Tesseract OCR", path_label="Tesseract executable path")
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)
